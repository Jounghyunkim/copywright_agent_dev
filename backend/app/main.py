import os
import uuid
import json as json_module
from datetime import datetime, timezone
from contextlib import asynccontextmanager
from typing import Optional
from dotenv import load_dotenv
import tempfile
from fastapi import FastAPI, File, Form, HTTPException, Depends, UploadFile
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from .schemas import (
    CampaignBrief, AnalyzeRequest, AnalysisResponse, ChatRequest, ChatResponse,
    GenerateBriefRequest, GenerateBriefResponse,
    StrategicMessageRequest, StrategicMessageResponse,
    GenerateCopyRequest, GenerateCopyResponse,
    ReviewRequest, ReviewSessionResponse, ReviewHistoryResponse,
    CorrectionRequest, CorrectionResponse,
    CustomSkillCreate, CustomSkillUpdate, SkillResponse,
    SkillDraftRequest,
    CampaignSaveRequest,
    MessageMatrixSheetsResponse, MessageMatrixParseResponse,
    MessageMatrixProduct, MessageMatrixCategory, MessageMatrixUSP,
    ExtractFilesResponse, ExtractedFile,
)
from .graph import app_graph
from .database import init_db, get_db, async_session
from .models import ReviewSession, ReviewResult, Campaign, AdminUser
from sqlalchemy import func
from .skills.runner import run_review
from .skills.catalog import get_all_skills, get_skillmd_skills
from .auth.redis_store import get_redis, close_redis
from .auth.middleware import AuthContextMiddleware, AuthContext, require_auth
from .auth.routes import router as auth_router
from .auth.admin_routes import router as admin_router
from .auth.stats_routes import router as stats_router
from .knowledge_routes import router as knowledge_router

load_dotenv(dotenv_path='.env')


@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    print("Database tables initialized.")
    # Redis 연결 사전 검증
    try:
        redis = await get_redis()
        await redis.ping()
        print("Redis connected.")
    except Exception as e:
        print(f"WARNING: Redis not available — auth will fail: {e}")
    # 초기 관리자 부트스트랩 (테이블 비어있을 때만)
    initial_admins = os.getenv("INITIAL_ADMIN_USER_IDS", "")
    if initial_admins:
        async with async_session() as db:
            count = (await db.execute(func.count(AdminUser.id))).scalar() or 0
            if count == 0:
                for uid in initial_admins.split(","):
                    uid = uid.strip()
                    if uid:
                        db.add(AdminUser(
                            user_id=uid,
                            display_name=uid,
                            added_by="system_bootstrap",
                        ))
                await db.commit()
                print(f"Bootstrapped initial admins: {initial_admins}")
    yield
    await close_redis()


app = FastAPI(title="Copywrite Agent API", lifespan=lifespan)

# --- CORS Middleware ---
origins = [
    "http://localhost",
    "http://localhost:5173",
    "http://localhost:5001", # Added new frontend port
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Auth Middleware (Redis session) ---
# Note: must be added AFTER CORS middleware (starlette processes in reverse order)
import asyncio as _auth_asyncio

async def _get_redis_for_middleware():
    return await get_redis()

class _RedisSessionAdapter:
    """Adapter to make Redis client compatible with middleware's get(key) interface."""
    async def get(self, key: str):
        redis = await get_redis()
        return await redis.get(key)

app.add_middleware(
    AuthContextMiddleware,
    session_store=_RedisSessionAdapter(),
    system_name="copywriting-agent",
    env=os.getenv("AUTH_ENV", "dev"),
    public_paths={
        "/",
        "/health",
        "/docs",
        "/openapi.json",
        "/auth/login",
        "/auth/logout",
    },
)

# --- Auth Routes ---
app.include_router(auth_router)
app.include_router(admin_router)
app.include_router(stats_router)
app.include_router(knowledge_router)


@app.get("/")
async def root():
    return {"message": "Copywrite Agent Backend API is running"}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.post("/api/v1/campaigns/analyze")
async def analyze_campaign(req: AnalyzeRequest):
    print(f"Received analysis request for: {req.projectName}")

    def _sse(data: dict) -> str:
        return f"data: {json_module.dumps(data, ensure_ascii=False)}\n\n"

    async def event_stream():
        yield _sse({"type": "progress", "message": f"Received analysis request for: {req.projectName}"})

        brief_dict = req.dict(exclude={"message_matrix", "locale"})
        inputs = {"brief": brief_dict, "message_matrix": req.message_matrix or {}, "locale": req.locale or "ko"}
        completed_parallel = set()
        final_report = None

        try:
            async for event in app_graph.astream(inputs, stream_mode="updates"):
                for node_name, node_output in event.items():
                    if node_name == "query_planner":
                        queries = node_output.get("search_queries", [])
                        yield _sse({"type": "progress", "message": "--- NODE 1: QUERY PLANNER ---"})
                        q_list = ", ".join(f"'{q}'" for q in queries)
                        yield _sse({"type": "progress", "message": f"Generated {len(queries)} search queries: [{q_list}]"})
                        yield _sse({"type": "progress", "message": "--- NODE 2a: WEB SEARCH ---"})
                        yield _sse({"type": "progress", "message": "--- NODE 2b: ENHANCED RAG ---"})

                    elif node_name == "web_search":
                        results = node_output.get("web_results", [])
                        yield _sse({"type": "progress", "message": f"Web search returned {len(results)} unique results."})
                        completed_parallel.add("web_search")
                        if completed_parallel >= {"web_search", "enhanced_rag"}:
                            yield _sse({"type": "progress", "message": "--- NODE 3: SYNTHESIZER ---"})
                            yield _sse({"type": "progress", "message": "Invoking synthesizer LLM..."})

                    elif node_name == "enhanced_rag":
                        results = node_output.get("rag_results", [])
                        yield _sse({"type": "progress", "message": f"RAG retrieved {len(results)} unique documents."})
                        completed_parallel.add("enhanced_rag")
                        if completed_parallel >= {"web_search", "enhanced_rag"}:
                            yield _sse({"type": "progress", "message": "--- NODE 3: SYNTHESIZER ---"})
                            yield _sse({"type": "progress", "message": "Invoking synthesizer LLM..."})

                    elif node_name == "synthesizer":
                        final_report = node_output.get("analysis_report", {})

            yield _sse({"type": "result", "status": "success", "data": final_report})

        except FileNotFoundError as e:
            print(f"FAISS index not found: {e}")
            yield _sse({"type": "error", "message": "Knowledge base index not found. Please run the data ingestion script first."})
        except Exception as e:
            print(f"Analysis failed: {e}")
            yield _sse({"type": "error", "message": f"Analysis failed: {str(e)}"})

        yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.post("/api/v1/campaigns/strategic-message", response_model=StrategicMessageResponse)
async def extract_strategic_message(request: StrategicMessageRequest):
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import HumanMessage, SystemMessage
    from langchain_core.output_parsers import JsonOutputParser
    import json

    print(f"Extracting strategic message from analysis report...")

    llm = AzureChatOpenAI(
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        temperature=0.7,
    )

    system_prompt = """You are a senior Brand Strategist at LG Electronics, specializing in strategic message architecture.

Given a Campaign Brief and its Market Analyst Report, extract and synthesize the **Strategic Message** — the core communication strategy that will guide all downstream copy generation.

Your output must be a JSON object with these exact keys:

{
  "coreMessage": "The single, overarching strategic message (1-2 sentences) that captures the essence of what this campaign must communicate. It should bridge the brand's value proposition with the consumer's emotional need.",
  "messagePillars": [
    {
      "title": "Pillar name (e.g., 'Emotional Connection', 'Functional Excellence')",
      "description": "How this pillar supports the core message and connects to the target persona's needs"
    }
  ],
  "emotionalHook": "The primary emotional trigger that will capture attention and drive engagement — derived from the Emotional JTBD and Cultural Tension insights",
  "toneDirection": {
    "primary": "The dominant tone (e.g., 'Warm yet authoritative', 'Aspirational but grounded')",
    "avoid": "Tones or approaches to avoid based on competitive analysis and brand fit",
    "voiceCharacter": "If this brand were a person speaking, describe their character in one sentence"
  },
  "keyPhrases": ["phrase1", "phrase2", "phrase3", "...up to 8 key phrases that should appear in or inspire the final copy"]
}

IMPORTANT:
- The coreMessage must directly stem from the emotionalJTBD and categoryNarrative shift
- messagePillars should be 2-4 items, each grounded in the analysis report findings
- keyPhrases should include both recommended keywords and newly synthesized phrases
- Everything must be aligned with the brand fit assessment and copy implications
- Write ALL output in {output_language}. This is the user's chosen display language."""

    locale = request.locale or "ko"
    output_language = {"en": "English", "ko": "Korean", "de": "German"}.get(locale, "Korean")
    system_prompt = system_prompt.replace("{output_language}", output_language)

    try:
        parser = JsonOutputParser()
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"## Campaign Brief\n```json\n{json.dumps(request.brief, ensure_ascii=False, indent=2)}\n```\n\n## Market Analyst Report\n```json\n{json.dumps(request.analysisReport, ensure_ascii=False, indent=2)}\n```\n\nPlease extract the Strategic Message based on the above inputs."),
        ]
        response = await llm.ainvoke(messages)
        data = parser.parse(response.content)
        return {"status": "success", "data": data}
    except Exception as e:
        print(f"Strategic message extraction failed: {e}")
        raise HTTPException(status_code=500, detail=f"Strategic message extraction failed: {str(e)}")


@app.post("/api/v1/campaigns/generate-copy", response_model=GenerateCopyResponse)
async def generate_copy(request: GenerateCopyRequest):
    """DeepAgent 기반 카피 생성 — SKILL.md 스킬 자동 선택 + 컨텍스트 주입"""
    from .skills.deep_agent import DeepAgentExecutor

    config = request.config
    countries_str = ", ".join(config.countries)
    print(f"[DeepAgent] Generating copy for countries: {countries_str}")

    try:
        agent = DeepAgentExecutor()
        result = await agent.generate_copy(
            brief=request.brief,
            analysis_report=request.analysisReport,
            strategic_message=request.strategicMessage,
            config={
                "countries": config.countries,
                "ageGroups": config.ageGroups,
                "personas": config.personas,
                "skillsets": config.skillsets,
                "copyCount": config.copyCount,
            },
            persona_skill=config.writerPersona or None,
        )

        print(f"[DeepAgent] Selected skills: {result['selected_skills']}")
        print(f"[DeepAgent] Skill reviews: {result.get('skill_reviews', {})}")
        print(f"[DeepAgent] Elapsed: {result['elapsed_ms']}ms")
        if result.get("diagnostics"):
            total_warns = sum(
                len(d.get("warnings", [])) for country in result["diagnostics"] for d in country
            )
            print(f"[DeepAgent] Filter warnings: {total_warns}")

        resp: dict = {"status": "success", "data": result["copies"]}
        if result.get("diagnostics"):
            resp["diagnostics"] = result["diagnostics"]
        return resp
    except Exception as e:
        print(f"Copy generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Copy generation failed: {str(e)}")


@app.get("/api/v1/culture-profiles")
async def list_culture_profiles():
    """사용 가능한 국가별 문화 프로필 목록"""
    from .skills.loader import SkillLoader
    loader = SkillLoader()
    profiles = loader.list_culture_profiles()
    return {
        "status": "success",
        "data": [
            {
                "id": p["name"],
                "description": p["description"],
                "country_code": p["name"].replace("culture-", "").upper(),
            }
            for p in profiles
        ],
    }


@app.get("/api/v1/skills/catalog")
async def get_full_skill_catalog():
    """전체 스킬 카탈로그 (SKILL.md 포함, 카테고리/태그별 분류)"""
    from .skills.catalog import get_all_skills
    skills = get_all_skills()
    # 카테고리별 분류
    by_category: dict[str, list] = {}
    for s in skills:
        cat = s.get("category", "other")
        by_category.setdefault(cat, []).append(s)
    return {
        "status": "success",
        "total": len(skills),
        "by_category": by_category,
        "data": skills,
    }


@app.post("/api/v1/campaigns/generate-copy-candidates")
async def generate_copy_candidates(request: GenerateCopyRequest):
    """페르소나 기반 다중 후보 생성 — 각 페르소나별 카피 후보를 반환"""
    from .skills.deep_agent import DeepAgentExecutor

    config = request.config
    print(f"[DeepAgent] Generating persona candidates for: {', '.join(config.countries)}")

    try:
        agent = DeepAgentExecutor()
        result = await agent.generate_copy_with_personas(
            brief=request.brief,
            analysis_report=request.analysisReport,
            strategic_message=request.strategicMessage,
            config={
                "countries": config.countries,
                "ageGroups": config.ageGroups,
                "personas": config.personas,
                "skillsets": config.skillsets,
                "copyCount": config.copyCount,
            },
        )

        print(f"[DeepAgent] Personas used: {[p['id'] for p in result.get('selected_personas', [])]}")
        print(f"[DeepAgent] Candidates generated: {len(result.get('candidates', []))}")
        print(f"[DeepAgent] Elapsed: {result['elapsed_ms']}ms")

        return {"status": "success", "data": result}
    except Exception as e:
        print(f"Persona candidate generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Candidate generation failed: {str(e)}")


@app.get("/api/v1/personas")
async def list_personas():
    """사용 가능한 AI Writer 페르소나 목록"""
    from .skills.creative_personas import get_creative_personas
    personas = get_creative_personas()
    return {
        "status": "success",
        "data": [
            {
                "id": p["id"],
                "name": p["name"],
                "avatar": p.get("avatar", ""),
                "color": p.get("color", "#9ca3af"),
                "tags": p.get("tags", []),
                "temperature": p.get("temperature", 0.9),
                "description": p.get("description", ""),
                "style_highlights": p.get("style_highlights", []),
            }
            for p in personas
        ],
    }


@app.post("/api/v1/campaigns/generate-brief", response_model=GenerateBriefResponse)
async def generate_brief(request: GenerateBriefRequest):
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import HumanMessage, SystemMessage
    from langchain_core.output_parsers import JsonOutputParser

    print(f"Generating brief draft for: {request.projectName}")

    llm = AzureChatOpenAI(
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        temperature=0.7,
    )

    # Load objective generation guide from file
    guide_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data", "brief_objective_gen_guide.md")
    try:
        with open(guide_path, "r", encoding="utf-8") as f:
            objective_guide = f.read()
    except FileNotFoundError:
        objective_guide = ""
        print(f"WARNING: Objective guide not found at {guide_path}")

    brief_guide = f"""You are a senior Brand Strategist with 15+ years of experience in global campaign planning for LG Electronics.
You are given a Project Name and a Project Context that describes WHY this campaign is needed. Use these two inputs to generate the REMAINING fields of an LG standard campaign brief.

Your output must be deeply informed by the Project Context — the objectives, audience, messaging, and strategy should all directly stem from the stated context and reasons for the campaign.

## Objective 작성 지침 (매우 중요 — 반드시 아래 가이드를 따르세요)

{objective_guide}

## 나머지 항목 작성 지침

나머지 항목들도 Project Context에서 논리적으로 파생되어야 합니다.

## 출력 형식

Return ONLY a valid JSON object with these exact keys (all values as strings in Korean):
{{
  "objectiveCommercial": "위 Objective 가이드의 Commercial 지침에 따라 1~2개 항목을 작성. 각 항목은 '~하는 것입니다' 형태의 완결된 문장.",
  "objectiveBehavior": "위 Objective 가이드의 Behavior 지침에 따라 1~2개 항목을 작성. 타깃의 동기를 반영한 완결된 문장.",
  "objectiveAttitudinal": "위 Objective 가이드의 Attitudinal 지침에 따라 1~2개 항목을 작성. Context의 고유 키워드를 활용한 완결된 문장.",
  "audience": "Primary/Secondary 타겟을 Project Context의 제품/시장에 맞춰 구체적으로 정의 — 인구통계, 라이프스타일, 구매 행동, 페인 포인트 포함",
  "keyMessage": "Project Context의 핵심 가치를 소비자 관점의 Benefit으로 변환한 1-2문장 (기능 나열 아닌 감정적 가치, Life's Good 연결)",
  "proofPoints": "Key Message 뒷받침 근거 1~3개 — Project Context에서 언급된 기술적 우위, 수상, 실적 등을 구체적 출처와 함께 서술",
  "mandatories": "Project Context의 시장/채널 특성에 맞는 매체 믹스, 제작물 스펙, 브랜드 가이드라인 준수 사항",
  "budget": "캠페인 규모에 적합한 총 예산 및 매체별 배분 비중 제안",
  "marketNeeds": "Project Context에서 파악한 타겟 국가/지역, 언어 variation, 문화적 고려사항, 현지 경쟁 환경",
  "timing": "Project Context의 런칭/시즌 정보에 맞춘 집행 기간, 페이즈 구분 (티징→런칭→서스테인), 제작 마감일"
}}

IMPORTANT: Every field must logically connect back to the Project Context. Do not generate generic content — make it specific to THIS project."""

    try:
        parser = JsonOutputParser()
        messages = [
            SystemMessage(content=brief_guide),
            HumanMessage(content=f"프로젝트명: {request.projectName}\n\nProject Context:\n{request.projectContext}\n\n위 프로젝트의 배경과 맥락을 바탕으로, 나머지 브리프 항목을 JSON으로 작성해 주세요."),
        ]
        response = await llm.ainvoke(messages)
        data = parser.parse(response.content)
        return {"status": "success", "data": data}
    except Exception as e:
        print(f"Brief generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Brief generation failed: {str(e)}")


@app.post("/api/v1/campaigns/chat", response_model=ChatResponse)
async def chat_with_agent(request: ChatRequest):
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

    llm = AzureChatOpenAI(
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        temperature=0.7,
    )

    # Build step-aware system prompt with available context
    step = request.currentStep or 1
    ctx = request.context or {}

    base_instruction = (
        "You are a professional LG marketing copywriting expert and campaign strategist. "
        "Keep answers concise, practical, and actionable. Respond in the same language as the user.\n\n"
    )

    step_instructions = {
        1: "The user is currently writing a campaign brief. Help them fill out fields like target audience, key message, proof points, tone & manner, and market needs. Provide specific, actionable advice.",
        2: "The user is reviewing the Market Analyst Report generated from their brief. Help them interpret the analysis results, discuss persona insights, brand fit scores, competitive keywords, and emotional JTBD. Suggest whether to approve or request modifications.",
        3: "The user is working on the Strategic Message step. Help them refine the Core Message and Message Pillars. Discuss how the strategic direction aligns with the brief and analysis insights.",
        4: "The user is in the Copy Generation step. Help them evaluate generated copy variants, discuss headline/subheadline effectiveness, CTA strength, and cultural appropriateness for each target market.",
        5: "The user is in the Review step where generated copies are being validated by skill-based checks. Help them interpret review results, understand flagged issues, and decide on final copy selections.",
    }

    prompt_parts = [base_instruction, step_instructions.get(step, step_instructions[1])]

    # Attach available artifacts as reference context
    if ctx.get("brief"):
        prompt_parts.append(f"\n\n=== Campaign Brief ===\n{json_module.dumps(ctx['brief'], ensure_ascii=False, indent=2)}")
    if ctx.get("analysisReport"):
        prompt_parts.append(f"\n\n=== Market Analyst Report ===\n{json_module.dumps(ctx['analysisReport'], ensure_ascii=False, indent=2)}")
    if ctx.get("strategicMessage"):
        prompt_parts.append(f"\n\n=== Strategic Message ===\n{json_module.dumps(ctx['strategicMessage'], ensure_ascii=False, indent=2)}")
    if ctx.get("copyResults"):
        # Send summary to avoid token overflow
        copy_summary = []
        for c in ctx["copyResults"][:10]:
            copy_summary.append({k: v for k, v in c.items() if k in ("countryCode", "headline", "subheadline", "cta")})
        prompt_parts.append(f"\n\n=== Generated Copies (summary) ===\n{json_module.dumps(copy_summary, ensure_ascii=False, indent=2)}")
    if ctx.get("reviewResults"):
        prompt_parts.append(f"\n\n=== Review Results ===\n{json_module.dumps(ctx['reviewResults'], ensure_ascii=False, indent=2)}")

    system_prompt = SystemMessage(content="".join(prompt_parts))

    history = [system_prompt]
    for msg in request.messages:
        if msg.role == "user":
            # Partition attachments: text → inline into prompt body;
            # image → separate content block for vision-capable models.
            body = msg.content
            image_blocks: list[dict] = []
            if msg.attachments:
                for att in msg.attachments:
                    if att.kind == "image" and att.image_url:
                        image_blocks.append({
                            "type": "image_url",
                            "image_url": {"url": att.image_url, "detail": "auto"},
                        })
                        # Add a brief text hint so the model knows the filename
                        body += f"\n\n[첨부 이미지: {att.filename}]"
                    else:
                        att_text = att.text or ""
                        trunc_note = " (truncated)" if att.truncated else ""
                        body += (
                            f"\n\n=== Attached File: {att.filename}{trunc_note} ===\n"
                            f"{att_text}"
                        )

            if image_blocks:
                # Multi-modal content list — requires a vision-capable deployment
                # (e.g. gpt-4o, gpt-4-turbo with vision).
                content_blocks: list[dict] = [{"type": "text", "text": body}]
                content_blocks.extend(image_blocks)
                history.append(HumanMessage(content=content_blocks))
            else:
                history.append(HumanMessage(content=body))
        elif msg.role == "assistant":
            history.append(AIMessage(content=msg.content))

    try:
        response = await llm.ainvoke(history)
        return {"reply": response.content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# --- Chat attachment text extraction ---

MAX_CHAT_FILE_BYTES = 5 * 1024 * 1024      # 5 MB per file
MAX_CHAT_FILE_TEXT_CHARS = 40_000          # truncate extracted text per file


def _extract_file_text(filename: str, raw: bytes) -> tuple[str, bool, Optional[str]]:
    """Return (text, truncated, error) for a single uploaded file.

    Supported: txt, md, csv, json, log (decoded as UTF-8 with replacement),
    xlsx (openpyxl), pdf (pypdf), docx (python-docx).
    Unsupported types return error and empty text.
    """
    import pathlib as _pathlib
    import io
    import csv as _csv

    ext = _pathlib.Path(filename).suffix.lower().lstrip(".")
    text = ""
    error: Optional[str] = None

    try:
        if ext in {"txt", "md", "markdown", "log", "json", "csv", "tsv", ""}:
            text = raw.decode("utf-8", errors="replace")
        elif ext == "pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(raw))
                parts: list[str] = []
                for page in reader.pages:
                    try:
                        parts.append(page.extract_text() or "")
                    except Exception:
                        parts.append("")
                text = "\n".join(parts)
            except Exception as e:
                error = f"PDF parse failed: {e}"
        elif ext == "docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(raw))
                lines = [p.text for p in doc.paragraphs if p.text]
                for table in doc.tables:
                    for row in table.rows:
                        line = "\t".join(cell.text for cell in row.cells)
                        if line.strip():
                            lines.append(line)
                text = "\n".join(lines)
            except Exception as e:
                error = f"DOCX parse failed: {e}"
        elif ext in {"xlsx", "xlsm"}:
            try:
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
                blocks: list[str] = []
                for ws in wb.worksheets:
                    blocks.append(f"[sheet: {ws.title}]")
                    for row in ws.iter_rows(values_only=True):
                        cells = ["" if v is None else str(v) for v in row]
                        if any(c.strip() for c in cells):
                            blocks.append("\t".join(cells))
                text = "\n".join(blocks)
            except Exception as e:
                error = f"XLSX parse failed: {e}"
        else:
            error = f"unsupported file type: .{ext or 'unknown'}"
    except Exception as e:
        error = f"extraction failed: {e}"

    truncated = False
    if len(text) > MAX_CHAT_FILE_TEXT_CHARS:
        text = text[:MAX_CHAT_FILE_TEXT_CHARS]
        truncated = True

    return text, truncated, error


@app.post("/api/v1/chat/extract-text", response_model=ExtractFilesResponse)
async def extract_chat_attachments(files: list[UploadFile] = File(...)):
    """Extract plain-text from one or more uploaded files for AI chat context.

    Supported: txt/md/csv/json/log, pdf, docx, xlsx.
    Files larger than 5 MB are rejected. Extracted text is capped at 40k chars.
    """
    if not files:
        raise HTTPException(status_code=400, detail="no files uploaded")

    out: list[ExtractedFile] = []
    for f in files:
        raw = await f.read()
        size = len(raw)
        if size > MAX_CHAT_FILE_BYTES:
            out.append(ExtractedFile(
                filename=f.filename or "unknown",
                text="",
                size=size,
                truncated=False,
                error=f"file too large ({size} bytes, max {MAX_CHAT_FILE_BYTES})",
            ))
            continue

        text, truncated, error = _extract_file_text(f.filename or "unknown", raw)
        out.append(ExtractedFile(
            filename=f.filename or "unknown",
            text=text,
            size=size,
            truncated=truncated,
            error=error,
        ))

    return {"files": out}


# ============================================================
# Review Endpoints
# ============================================================

@app.post("/api/v1/campaigns/review")
async def submit_review(request: ReviewRequest):
    """Review 실행 — SSE 스트림으로 스킬별 진행률 + 최종 결과 반환"""
    project_name = request.brief.get("projectName", "unknown")
    print(f"Review requested for: {project_name}, skills: {request.enabledSkills}")

    def _sse(data: dict) -> str:
        return f"data: {json_module.dumps(data, ensure_ascii=False)}\n\n"

    # run_review 는 asyncio.gather로 전체를 모아 반환하므로, 대기 중 SSE가 끊길 수 있음.
    # 대신 asyncio.Queue를 사용해 태스크 완료 즉시 SSE로 스트리밍.
    import asyncio as _aio
    from .skills.runner import _get_skillmd_ids, _build_copy_text, _run_single
    from .skills.custom import get_custom_skill

    async def event_stream():
        async with async_session() as db:
            # 1. 세션 생성
            session = ReviewSession(
                project_name=project_name,
                brief_snapshot=request.brief,
                analysis_snapshot=request.analysisReport,
                strategic_message_snapshot=request.strategicMessage,
                selected_copies=[c.model_dump() for c in request.selectedCopies],
                enabled_skills=request.enabledSkills,
                status="running",
            )
            db.add(session)
            await db.commit()
            await db.refresh(session)
            session_id = str(session.id)

            yield _sse({"type": "review_started", "sessionId": session_id, "skills": request.enabledSkills})

            try:
                # 2. 태스크 준비
                copies_for_runner = [c.model_dump() for c in request.selectedCopies]
                skillmd_ids = _get_skillmd_ids()
                custom_templates = {}
                for sid in request.enabledSkills:
                    if sid not in skillmd_ids:
                        skill_data = get_custom_skill(sid)
                        if skill_data and skill_data.get("is_active", True):
                            custom_templates[sid] = skill_data["prompt_template"]

                # 3. Queue 기반 실시간 스트리밍
                result_queue: _aio.Queue = _aio.Queue()
                total_tasks = 0

                async def _wrapped(skill_id, skill_type, copy_key, copy_text, context, prompt_template):
                    raw = await _run_single(skill_id, skill_type, copy_text, context, prompt_template)
                    entry = {"skill_id": skill_id, "skill_type": skill_type, "target_copy_key": copy_key, **raw}
                    await result_queue.put(entry)

                tasks = []
                for skill_id in request.enabledSkills:
                    if skill_id in skillmd_ids:
                        skill_type, prompt_template = "skillmd", None
                    elif skill_id in custom_templates:
                        skill_type, prompt_template = "custom", custom_templates[skill_id]
                    else:
                        continue
                    for copy_entry in copies_for_runner:
                        copy_key = copy_entry["key"]
                        copy_data = copy_entry.get("copyData") or copy_entry.get("copy", {})
                        copy_text = _build_copy_text(copy_data)
                        context = {
                            "brief_summary": request.brief,
                            "analysis_context": request.analysisReport,
                            "strategic_message": request.strategicMessage,
                            "country_code": copy_entry.get("countryCode", ""),
                        }
                        tasks.append(_wrapped(skill_id, skill_type, copy_key, copy_text, context, prompt_template))
                        total_tasks += 1

                # 모든 태스크를 백그라운드로 시작
                gather_task = _aio.ensure_future(_aio.gather(*tasks, return_exceptions=True))

                # 4. 완료되는 순서대로 SSE 전송
                results_collected = []
                for _ in range(total_tasks):
                    r = await result_queue.get()
                    results_collected.append(r)
                    # DB에 저장
                    db_result = ReviewResult(
                        session_id=session.id,
                        skill_id=r["skill_id"],
                        skill_type=r["skill_type"],
                        target_copy_key=r["target_copy_key"],
                        passed=r["passed"],
                        score=r["score"],
                        findings=r.get("weaknesses", r.get("findings", [])),
                        suggestions=r.get("improvements", r.get("suggestions", [])),
                        raw_llm_response=r.get("raw_llm_response"),
                        execution_ms=r.get("execution_ms", 0),
                    )
                    db.add(db_result)
                    yield _sse({
                        "type": "skill_completed",
                        "skillId": r["skill_id"],
                        "skillType": r["skill_type"],
                        "targetCopyKey": r["target_copy_key"],
                        "passed": r["passed"],
                        "score": r["score"],
                        "strengths": r.get("strengths", []),
                        "weaknesses": r.get("weaknesses", r.get("findings", [])),
                        "improvements": r.get("improvements", r.get("suggestions", [])),
                        "executionMs": r.get("execution_ms", 0),
                    })

                # gather 대기 (이미 완료되었을 것이지만 예외 처리용)
                await gather_task

                # 5. 세션 완료
                session.status = "completed"
                session.completed_at = datetime.now(timezone.utc)
                await db.commit()

                # 6. 요약
                total = len(results_collected)
                passed = sum(1 for r in results_collected if r["passed"])
                avg_score = round(sum(r["score"] for r in results_collected) / total, 1) if total else 0
                yield _sse({
                    "type": "review_done",
                    "sessionId": session_id,
                    "summary": {
                        "total": total,
                        "passed": passed,
                        "failed": total - passed,
                        "avgScore": avg_score,
                    },
                })

            except Exception as e:
                print(f"Review failed: {e}")
                import traceback; traceback.print_exc()
                session.status = "failed"
                try:
                    await db.commit()
                except Exception:
                    pass
                yield _sse({"type": "error", "message": f"Review failed: {str(e)}"})

        yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.get("/api/v1/campaigns/review/history")
async def get_review_history(project_name: str = None, db: AsyncSession = Depends(get_db)):
    """프로젝트별 리뷰 이력 조회"""
    query = select(ReviewSession).order_by(ReviewSession.created_at.desc())
    if project_name:
        query = query.where(ReviewSession.project_name == project_name)
    query = query.limit(50)

    result = await db.execute(query)
    sessions = [
        {
            "id": str(s.id),
            "projectName": s.project_name,
            "status": s.status,
            "enabledSkills": s.enabled_skills,
            "createdAt": s.created_at.isoformat() if s.created_at else None,
            "completedAt": s.completed_at.isoformat() if s.completed_at else None,
        }
        for s in result.scalars().all()
    ]
    return {"sessions": sessions}


@app.get("/api/v1/campaigns/review/{session_id}")
async def get_review_session(session_id: str, db: AsyncSession = Depends(get_db)):
    """리뷰 세션 결과 조회"""
    try:
        uid = uuid.UUID(session_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid session ID format")

    result = await db.execute(
        select(ReviewSession).where(ReviewSession.id == uid)
    )
    session = result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Review session not found")

    # 결과 로딩
    results_q = await db.execute(
        select(ReviewResult).where(ReviewResult.session_id == uid)
    )
    results = [
        {
            "id": str(r.id),
            "skillId": r.skill_id,
            "skillType": r.skill_type,
            "targetCopyKey": r.target_copy_key,
            "passed": r.passed,
            "score": r.score,
            "findings": r.findings,
            "suggestions": r.suggestions,
            "executionMs": r.execution_ms,
            "createdAt": r.created_at.isoformat() if r.created_at else None,
        }
        for r in results_q.scalars().all()
    ]

    return {
        "id": str(session.id),
        "projectName": session.project_name,
        "status": session.status,
        "enabledSkills": session.enabled_skills,
        "createdAt": session.created_at.isoformat() if session.created_at else None,
        "completedAt": session.completed_at.isoformat() if session.completed_at else None,
        "results": results,
    }


# ============================================================
# Copy Correction Endpoint
# ============================================================

@app.post("/api/v1/campaigns/correct", response_model=CorrectionResponse)
async def correct_copy(request: CorrectionRequest):
    """선택된 보완사항을 기반으로 카피를 보정하여 반환."""
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import HumanMessage, SystemMessage

    llm = AzureChatOpenAI(
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        temperature=0.4,
    )

    copy = request.copyData
    improvements_text = "\n".join(
        f"- [{imp.skillId}] {imp.text}" for imp in request.improvements
    )

    system_prompt = SystemMessage(content=(
        "You are an expert LG brand copywriter. You will receive an existing marketing copy "
        "(headline, subheadline, bodyCopy, cta) and a list of review improvement suggestions "
        "from quality-check skills.\n\n"
        "Your task is to revise the copy to address ALL the improvement suggestions.\n\n"
        "CRITICAL RULES:\n"
        "1. Apply improvements to EVERY relevant part — headline, subheadline, bodyCopy, AND cta.\n"
        "2. The HEADLINE is the most visible element. If an improvement relates to tone, brand compliance, "
        "   keyword usage, or messaging clarity, the headline MUST be updated accordingly.\n"
        "3. Do NOT leave any field unchanged if the improvement logically applies to it.\n"
        "4. Preserve the original language, tone intent, and approximate length.\n"
        "5. Return ONLY a JSON object with these exact keys: headline, subheadline, bodyCopy, cta.\n"
        "6. Do not include any explanation, markdown, or code fences — just the raw JSON object.\n"
        "7. Keep the same language as the original copy."
    ))

    user_prompt = HumanMessage(content=(
        f"=== Original Copy ===\n"
        f"Headline: {copy.get('headline', '')}\n"
        f"Subheadline: {copy.get('subheadline', '')}\n"
        f"Body Copy: {copy.get('bodyCopy', '')}\n"
        f"CTA: {copy.get('cta', '')}\n\n"
        f"=== Improvement Suggestions ===\n{improvements_text}\n\n"
        f"Revise ALL four fields (headline, subheadline, bodyCopy, cta) to address the improvements above. "
        f"Pay special attention to updating the HEADLINE — it must reflect the key improvements. "
        f"Return only a JSON object with all four keys."
    ))

    try:
        response = await llm.ainvoke([system_prompt, user_prompt])
        content = response.content.strip()
        # Strip markdown code fences if present
        if content.startswith("```"):
            content = content.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        result = json_module.loads(content)
        return CorrectionResponse(
            headline=result.get("headline", copy.get("headline", "")),
            subheadline=result.get("subheadline", copy.get("subheadline", "")),
            bodyCopy=result.get("bodyCopy", copy.get("bodyCopy", "")),
            cta=result.get("cta", copy.get("cta", "")),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Copy correction failed: {str(e)}")


# ============================================================
# Skill CRUD Endpoints
# ============================================================

@app.get("/api/v1/skills")
async def list_skills():
    """전체 스킬 목록 (빌트인 6개 + 커스텀)"""
    skills = get_all_skills()
    return {"skills": skills}


@app.post("/api/v1/skills", status_code=201)
async def create_skill(skill: CustomSkillCreate):
    """커스텀 스킬 등록 — skills/custom/ 폴더에 JSON 파일로 저장"""
    from .skills.custom import get_custom_skill, save_custom_skill

    skillmd_ids = {s["id"] for s in get_skillmd_skills()}
    if skill.id in skillmd_ids:
        raise HTTPException(status_code=409, detail=f"Skill ID '{skill.id}' conflicts with an existing SKILL.md skill")

    if get_custom_skill(skill.id):
        raise HTTPException(status_code=409, detail=f"Skill ID '{skill.id}' already exists")

    if skill.category not in ("validation", "generation", "analysis"):
        raise HTTPException(status_code=422, detail="category must be 'validation', 'generation', or 'analysis'")

    data = save_custom_skill({
        "id": skill.id,
        "label": skill.label,
        "description": skill.description,
        "category": skill.category,
        "prompt_template": skill.prompt_template,
        "reference_docs": skill.reference_docs,
        "output_schema": skill.output_schema,
    })

    return {
        "id": data["id"],
        "label": data["label"],
        "description": data["description"],
        "category": data["category"],
        "type": "custom",
        "editable": True,
    }


@app.get("/api/v1/skills/{skill_id}")
async def get_skill(skill_id: str):
    """단일 스킬 상세 조회"""
    from .skills.custom import get_custom_skill

    # SKILL.md 기반 스킬 검색
    for s in get_skillmd_skills():
        if s["id"] == skill_id:
            return s

    # 커스텀 스킬 검색
    skill = get_custom_skill(skill_id)
    if not skill:
        raise HTTPException(status_code=404, detail="Skill not found")

    return {
        "id": skill["id"],
        "label": skill["label"],
        "description": skill["description"],
        "category": skill["category"],
        "prompt_template": skill.get("prompt_template"),
        "reference_docs": skill.get("reference_docs"),
        "output_schema": skill.get("output_schema"),
        "is_active": skill.get("is_active", True),
        "type": "custom",
        "editable": True,
        "createdAt": skill.get("created_at"),
        "updatedAt": skill.get("updated_at"),
    }


@app.put("/api/v1/skills/{skill_id}")
async def update_skill(skill_id: str, updates: CustomSkillUpdate):
    """커스텀 스킬 수정 (SKILL.md 스킬은 수정 불가)"""
    from .skills.custom import get_custom_skill, save_custom_skill

    if skill_id in {s["id"] for s in get_skillmd_skills()}:
        raise HTTPException(status_code=403, detail="SKILL.md skills cannot be modified via API")

    skill = get_custom_skill(skill_id)
    if not skill:
        raise HTTPException(status_code=404, detail="Skill not found")

    update_data = updates.model_dump(exclude_unset=True)
    if "category" in update_data and update_data["category"] not in ("validation", "generation", "analysis"):
        raise HTTPException(status_code=422, detail="category must be 'validation', 'generation', or 'analysis'")

    skill.update(update_data)
    save_custom_skill(skill)

    return {
        "id": skill["id"],
        "label": skill["label"],
        "description": skill["description"],
        "category": skill["category"],
        "type": "custom",
        "editable": True,
    }


@app.delete("/api/v1/skills/{skill_id}")
async def delete_skill(skill_id: str):
    """커스텀 스킬 삭제 (SKILL.md 스킬 삭제 불가)"""
    from .skills.custom import delete_custom_skill

    if skill_id in {s["id"] for s in get_skillmd_skills()}:
        raise HTTPException(status_code=403, detail="SKILL.md skills cannot be deleted via API")

    if not delete_custom_skill(skill_id):
        raise HTTPException(status_code=404, detail="Skill not found")

    return {"status": "deleted", "id": skill_id}


# ============================================================
# Skill Draft Generation
# ============================================================

@app.post("/api/v1/skills/generate-draft")
async def generate_skill_draft(request: SkillDraftRequest):
    """사용자의 핵심 입력으로부터 스킬 전체 초안을 AI로 생성"""
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import HumanMessage, SystemMessage
    from langchain_core.output_parsers import JsonOutputParser
    import json

    llm = AzureChatOpenAI(
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        temperature=0.5,
    )

    system_prompt = """You are a senior QA & copywriting expert at LG Electronics.
Given a user's description of WHY they need a review skill and WHAT it should do, generate a complete skill definition.

Return a JSON object with these exact keys:

{
  "id": "kebab-case-skill-id (unique, descriptive, e.g. 'promo-legal-check')",
  "label": "Human-readable skill name in Korean (e.g. '프로모션 법적 검증')",
  "description": "1-2 sentence description of what this skill does, in Korean",
  "category": "validation",
  "prompt_template": "A detailed prompt template that an LLM will use to evaluate copywriting. The prompt should:\n  - Clearly state the evaluation criteria\n  - Reference {{copies}}, {{brief}}, {{analysis_report}}, {{strategic_message}} as template variables\n  - Include scoring rubric (0-100)\n  - Define pass/fail threshold\n  - Specify output format with: score, passed, strengths[], weaknesses[], improvements[]\n  - Be written in Korean for the marketing team\n  - Include the good/bad examples provided by the user as reference cases",
  "output_schema": {
    "type": "object",
    "properties": {
      "score": {"type": "number", "description": "0-100 점수"},
      "passed": {"type": "boolean"},
      "strengths": {"type": "array", "items": {"type": "string"}},
      "weaknesses": {"type": "array", "items": {"type": "string"}},
      "improvements": {"type": "array", "items": {"type": "string"}}
    }
  }
}

IMPORTANT:
- The prompt_template is the most critical field — it must be thorough, precise, and actionable
- category should always be "validation" for review skills
- The id must be kebab-case, 3-5 words, unique and descriptive
- Write everything in Korean except the id and JSON keys"""

    user_content = f"""## 스킬 이름
{request.name}

## 작성 목적
{request.purpose}

## 스킬 목적
{request.goal}"""

    if request.goodExample:
        user_content += f"\n\n## 좋은 예시\n{request.goodExample}"
    if request.badExample:
        user_content += f"\n\n## 나쁜 예시\n{request.badExample}"

    user_content += "\n\n위 정보를 기반으로 완전한 스킬 정의를 JSON으로 생성해 주세요."

    try:
        parser = JsonOutputParser()
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=user_content),
        ]
        response = await llm.ainvoke(messages)
        data = parser.parse(response.content)
        return {"draft": data}
    except Exception as e:
        print(f"Skill draft generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Skill draft generation failed: {str(e)}")


# ============================================================
# Campaign Save / Dashboard Endpoints
# ============================================================

def _derive_campaign_fields(request: CampaignSaveRequest) -> dict:
    """Save/Update 공통: 파생 필드 계산"""
    countries = list({r.get("countryCode", "") for r in (request.copyResults or []) if r.get("countryCode")})
    brand_fit = (request.analysisReport or {}).get("brandFit", {})
    brand_fit_score = brand_fit.get("score", 0) if isinstance(brand_fit, dict) else 0
    review_avg = round(request.reviewSummary["avgScore"]) if request.reviewSummary and request.reviewSummary.get("avgScore") else 0
    total_copies = sum(len(r.get("copies", [r])) for r in (request.copyResults or []))
    # 상태: step 5 완료 + reviewSummary가 있으면 completed, 아니면 draft
    status = "completed" if request.currentStep >= 5 and request.reviewSummary else "draft"
    return dict(
        target_countries=countries, brand_fit_score=brand_fit_score,
        review_avg_score=review_avg, total_copies=total_copies, status=status,
    )


@app.post("/api/v1/campaigns/save", status_code=201)
async def save_campaign(
    request: CampaignSaveRequest,
    db: AsyncSession = Depends(get_db),
    ctx: AuthContext = Depends(require_auth),
):
    """캠페인 저장 (신규 생성) — 각 단계 완료 시 자동 저장"""
    brief = request.brief
    project_name = brief.get("projectName", "Untitled")
    derived = _derive_campaign_fields(request)

    campaign = Campaign(
        project_name=project_name,
        brief=brief,
        analysis_report=request.analysisReport,
        strategic_message=request.strategicMessage,
        copy_results=request.copyResults,
        review_summary=request.reviewSummary,
        review_results=request.reviewResults,
        copy_candidates=request.copyCandidates,
        current_step=request.currentStep,
        created_by=ctx.user_id,
        **derived,
    )
    db.add(campaign)
    await db.commit()
    await db.refresh(campaign)

    return {
        "id": str(campaign.id),
        "projectName": campaign.project_name,
        "status": "saved",
        "currentStep": campaign.current_step,
    }


@app.get("/api/v1/campaigns/dashboard")
async def get_dashboard(
    db: AsyncSession = Depends(get_db),
    ctx: AuthContext = Depends(require_auth),
):
    """Dashboard 통계 + 최근 캠페인 목록"""
    is_admin = "admin" in ctx.roles
    query = select(Campaign).order_by(Campaign.created_at.desc()).limit(20)
    if not is_admin:
        query = query.where(Campaign.created_by == ctx.user_id)
    result = await db.execute(query)
    campaigns = result.scalars().all()

    # 통계 계산
    total_projects = len(campaigns)
    all_countries = set()
    brand_scores = []
    review_scores = []
    for c in campaigns:
        all_countries.update(c.target_countries or [])
        if c.brand_fit_score > 0:
            brand_scores.append(c.brand_fit_score)
        if c.review_avg_score > 0:
            review_scores.append(c.review_avg_score)

    avg_brand_score = round(sum(brand_scores) / len(brand_scores), 1) if brand_scores else 0
    avg_review_score = round(sum(review_scores) / len(review_scores), 1) if review_scores else 0
    target_regions = len(all_countries)

    campaign_list = [
        {
            "id": str(c.id),
            "title": c.project_name,
            "countries": c.target_countries or [],
            "date": c.created_at.strftime("%Y-%m-%d") if c.created_at else "",
            "brandFitScore": c.brand_fit_score,
            "reviewAvgScore": c.review_avg_score,
            "totalCopies": c.total_copies,
            "currentStep": c.current_step,
            "status": c.status,
            "summary": c.brief.get("projectContext", "")[:120] if isinstance(c.brief, dict) else "",
            "createdBy": c.created_by,
        }
        for c in campaigns
    ]

    return {
        "stats": {
            "totalProjects": total_projects,
            "avgBrandScore": avg_brand_score,
            "avgReviewScore": avg_review_score,
            "targetRegions": target_regions,
        },
        "campaigns": campaign_list,
    }


@app.get("/api/v1/campaigns/{campaign_id}")
async def get_campaign(
    campaign_id: str,
    db: AsyncSession = Depends(get_db),
    ctx: AuthContext = Depends(require_auth),
):
    """캠페인 상세 조회 — 전체 산출물 반환"""
    try:
        uid = uuid.UUID(campaign_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid campaign ID format")

    result = await db.execute(select(Campaign).where(Campaign.id == uid))
    campaign = result.scalar_one_or_none()
    if not campaign:
        raise HTTPException(status_code=404, detail="Campaign not found")

    is_admin = "admin" in ctx.roles
    if not is_admin and campaign.created_by and campaign.created_by != ctx.user_id:
        raise HTTPException(status_code=403, detail="이 캠페인에 대한 접근 권한이 없습니다.")

    return {
        "id": str(campaign.id),
        "projectName": campaign.project_name,
        "brief": campaign.brief,
        "analysisReport": campaign.analysis_report,
        "strategicMessage": campaign.strategic_message,
        "copyResults": campaign.copy_results,
        "reviewSummary": campaign.review_summary,
        "reviewResults": campaign.review_results,
        "copyCandidates": campaign.copy_candidates,
        "targetCountries": campaign.target_countries,
        "brandFitScore": campaign.brand_fit_score,
        "reviewAvgScore": campaign.review_avg_score,
        "totalCopies": campaign.total_copies,
        "currentStep": campaign.current_step,
        "status": campaign.status,
        "createdBy": campaign.created_by,
        "createdAt": campaign.created_at.isoformat() if campaign.created_at else None,
    }


@app.delete("/api/v1/campaigns/{campaign_id}")
async def delete_campaign(
    campaign_id: str,
    db: AsyncSession = Depends(get_db),
    ctx: AuthContext = Depends(require_auth),
):
    """캠페인 삭제"""
    try:
        uid = uuid.UUID(campaign_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid campaign ID format")

    result = await db.execute(select(Campaign).where(Campaign.id == uid))
    campaign = result.scalar_one_or_none()
    if not campaign:
        raise HTTPException(status_code=404, detail="Campaign not found")

    is_admin = "admin" in ctx.roles
    if not is_admin and campaign.created_by and campaign.created_by != ctx.user_id:
        raise HTTPException(status_code=403, detail="이 캠페인을 삭제할 권한이 없습니다.")

    await db.delete(campaign)
    await db.commit()
    return {"status": "deleted", "id": campaign_id}


@app.put("/api/v1/campaigns/{campaign_id}")
async def update_campaign(
    campaign_id: str,
    request: CampaignSaveRequest,
    db: AsyncSession = Depends(get_db),
    ctx: AuthContext = Depends(require_auth),
):
    """기존 캠페인 업데이트 — 단계 진행 시 자동 저장"""
    try:
        uid = uuid.UUID(campaign_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid campaign ID format")

    result = await db.execute(select(Campaign).where(Campaign.id == uid))
    campaign = result.scalar_one_or_none()
    if not campaign:
        raise HTTPException(status_code=404, detail="Campaign not found")

    is_admin = "admin" in ctx.roles
    if not is_admin and campaign.created_by and campaign.created_by != ctx.user_id:
        raise HTTPException(status_code=403, detail="이 캠페인을 수정할 권한이 없습니다.")

    brief = request.brief
    derived = _derive_campaign_fields(request)

    campaign.project_name = brief.get("projectName", campaign.project_name)
    campaign.brief = brief
    campaign.analysis_report = request.analysisReport
    campaign.strategic_message = request.strategicMessage
    campaign.copy_results = request.copyResults
    campaign.review_summary = request.reviewSummary
    campaign.review_results = request.reviewResults
    campaign.copy_candidates = request.copyCandidates
    campaign.current_step = request.currentStep
    campaign.target_countries = derived["target_countries"]
    campaign.brand_fit_score = derived["brand_fit_score"]
    campaign.review_avg_score = derived["review_avg_score"]
    campaign.total_copies = derived["total_copies"]
    campaign.status = derived["status"]

    await db.commit()
    await db.refresh(campaign)

    return {
        "id": str(campaign.id),
        "projectName": campaign.project_name,
        "status": "updated",
        "currentStep": campaign.current_step,
    }


# =============================================================
# Message Matrix — upload & parse
# =============================================================

def _product_info_to_dict(product) -> dict:
    """tools.message_matrix_parsing.ProductInfo dataclass → serializable dict."""
    return {
        "product_name": product.product_name,
        "sub_name": product.sub_name,
        "head_message": product.head_message,
        "description": product.description,
        "categories": [
            {
                "number": cat.number,
                "name": cat.name,
                "key_message": cat.key_message,
                "usps": [
                    {
                        "usp_no": u.usp_no,
                        "feature_name": u.feature_name,
                        "key_message_full": u.key_message_full,
                        "key_message_short": u.key_message_short,
                        "benefit_description": u.benefit_description,
                        "rtb": u.rtb,
                        "disclaimer": u.disclaimer,
                        "certification": u.certification,
                        "remark": u.remark,
                    }
                    for u in cat.usps
                ],
            }
            for cat in product.categories
        ],
    }


@app.post("/api/v1/message-matrix/sheets", response_model=MessageMatrixSheetsResponse)
async def message_matrix_sheets(file: UploadFile = File(...)):
    """xlsx 파일을 받아 시트 이름 목록을 반환."""
    import sys, pathlib
    tools_dir = str(pathlib.Path(__file__).resolve().parent.parent / "tools")
    if tools_dir not in sys.path:
        sys.path.insert(0, tools_dir)
    from message_matrix_parsing import get_sheet_names

    suffix = pathlib.Path(file.filename or "upload.xlsx").suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        sheets = get_sheet_names(tmp_path)
    finally:
        os.unlink(tmp_path)

    return {"sheets": sheets}


@app.post("/api/v1/message-matrix/parse", response_model=MessageMatrixParseResponse)
async def message_matrix_parse(
    file: UploadFile = File(...),
    sheets: str = Form(""),  # comma-separated sheet names, empty = all
):
    """xlsx 파일 + 선택 시트를 받아 파싱 결과를 반환."""
    import sys, pathlib
    tools_dir = str(pathlib.Path(__file__).resolve().parent.parent / "tools")
    if tools_dir not in sys.path:
        sys.path.insert(0, tools_dir)
    from message_matrix_parsing import parse_excel

    suffix = pathlib.Path(file.filename or "upload.xlsx").suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        sheet_list = [s.strip() for s in sheets.split(",") if s.strip()] or None
        products = parse_excel(tmp_path, sheet_names=sheet_list)
    finally:
        os.unlink(tmp_path)

    results = {name: _product_info_to_dict(p) for name, p in products.items()}
    return {"results": results}


@app.get("/api/v1/message-matrix/sample", response_model=MessageMatrixParseResponse)
async def message_matrix_sample():
    """테스트용: backend/tools/ 의 샘플 xlsx 를 파싱해 반환."""
    import sys, pathlib
    tools_dir = str(pathlib.Path(__file__).resolve().parent.parent / "tools")
    if tools_dir not in sys.path:
        sys.path.insert(0, tools_dir)
    from message_matrix_parsing import parse_excel

    sample_path = pathlib.Path(tools_dir) / "SAMPLE_(RAC)DUALCOOL AI Air_Message Matrix_v1.1 1.xlsx"
    if not sample_path.exists():
        raise HTTPException(status_code=404, detail="Sample file not found")

    products = parse_excel(str(sample_path))
    results = {name: _product_info_to_dict(p) for name, p in products.items()}
    return {"results": results}
